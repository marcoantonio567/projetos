from openpyxl import load_workbook
from docx import Document
from datetime import datetime

supplier_spreadsheet = load_workbook(r'C:\Users\marco\Desktop\update_repositorio\projetos_rpa\projetos\projeto1\fornecedores.xlsx')
sheet = supplier_spreadsheet['Sheet1']

for row in sheet.iter_rows(min_row=2):
    provider_name, address, city, state, CEP, phone, Email, material_type = row
    word_file = Document()
    word_file.add_heading('Contrato de prestação de serviço')

    text_content = f"""
Este contrato de prestação de serviços é feito entre {provider_name.value}, com address em {address.value}, 
{city.value}, {state.value}, CEP {CEP.value}, doravante denominado FORNECEDOR, e a empresa CONTRATANTE.

Pelo presente instrumento particular, as partes têm, entre si, justo e acordado o seguinte:

1. OBJETO DO CONTRATO
O FORNECEDOR compromete-se a fornecer à CONTRATANTE os serviços/material de acordo com as especificações acordadas, respeitando os padrões de qualidade e os prazos estipulados.

2. PRAZO
Este contrato tem prazo de vigência de 12 (doze) meses, iniciando-se na data de sua assinatura, podendo ser renovado conforme acordo entre as partes.

3. VALOR E FORMA DE PAGAMENTO
O valor dos serviços prstates será acordado conforme as demandas da CONTRATANTE e a capacity de entrega do FORNECEDOR. Os pagamentos serão realizados mensalmente, mediante apresentação de nota fiscal.

4. CONFIDENCIALIDADE
Todas as informações trocadas entre as partes durante a vigência deste contrato serão tratadas como confidenciais.

Para firmeza e como prova de assim haverem justo e contratado, as partes assinam o presente contrato em duas vias de igual teor e forma.

FORNECEDOR: {provider_name.value}
E-mail: {Email.value}

CONTRATANTE: Empresa_contratante
E-mail: Empresa_contratante@gmail.com

São Paulo ,{datetime.now().strftime('%d/%m/%Y')}
"""
    
    word_file.add_paragraph(text_content)
    
    word_file.save(f'C:\\Users\\marco\\Desktop\\update_repositorio\\projetos_rpa\\projetos\\projeto1\\contratos\\contrato_{provider_name.value}.docx')
